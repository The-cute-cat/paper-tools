"""将翻译后的论文导出为 DOCX / PDF 格式。

优先基于保留原 HTML 结构的 `.zh.html`（由 pipeline 在翻译完成后生成），
从而支持论文中更复杂的排版：

  * 表格合并单元格（colspan / rowspan）—— DOCX 用原生 cell.merge，
    PDF 按合并区域绘制一次；而 Markdown 只能降级为重复填充。
  * 内联文字颜色（原 HTML 上的 color 样式）—— DOCX 写 run.color，
    PDF 写 set_text_color。
  * 图片内嵌——Word / PDF 不支持"图片以链接存在"，这里会把本地图片
    或下载到的图片嵌入文档；Markdown 中图片只是链接，无法内嵌。

若只提供了 `.zh.md`（无 `.zh.html`），则退化为基于 Markdown 的降级导出。
"""

from __future__ import annotations

import re
import urllib.request
from html import escape, unescape
from pathlib import Path
from urllib.parse import urljoin

from bs4 import BeautifulSoup
from fpdf.enums import XPos, YPos

from ..logging_setup import get_logger

logger = get_logger()

# ──────────────────────────────────────────────
# 通用 Markdown 解析工具（降级路径）
# ──────────────────────────────────────────────

_BLOCK_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BLOCK_HR_RE = re.compile(r"^[-*_]{3,}\s*$")
_BLOCK_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
_BLOCK_TABLE_SEP_RE = re.compile(r"^\|[\s\-:|\s]+\|\s*$")
_BLOCK_TABLE_ROW_RE = re.compile(r"^\|.+\|\s*$")
_BLOCK_IMAGE_RE = re.compile(r"^!\[.*?\]\((.+?)\)\s*$")

# 内联格式
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_INLINE_MATH_RE = re.compile(r"\$(.+?)\$")
_LINK_RE = re.compile(r"\[([^\]]*)\]\(([^)]*)\)")

# 把 html 颜色名/hex 转 (r,g,b)
_COLOR_NAME_MAP = {
    "black": (0, 0, 0), "red": (255, 0, 0), "green": (0, 128, 0),
    "blue": (0, 0, 255), "gray": (128, 128, 128), "grey": (128, 128, 128),
    "orange": (255, 165, 0), "purple": (128, 0, 128), "brown": (165, 42, 42),
}


def _strip_inline(text: str) -> str:
    """去除 Markdown 内联标记，返回纯文本。"""
    t = text
    t = _BOLD_RE.sub(r"\1", t)
    t = _ITALIC_RE.sub(r"\1", t)
    t = _INLINE_CODE_RE.sub(r"\1", t)
    t = _LINK_RE.sub(r"\1", t)
    return t


def _parse_color(value: str | None) -> tuple[int, int, int] | None:
    """解析 CSS 颜色为 (r,g,b)；无法解析返回 None。"""
    if not value:
        return None
    v = value.strip().lower()
    if v in _COLOR_NAME_MAP:
        return _COLOR_NAME_MAP[v]
    m = re.match(r"#([0-9a-f]{2})([0-9a-f]{2})([0-9a-f]{2})", v)
    if m:
        return int(m.group(1), 16), int(m.group(2), 16), int(m.group(3), 16)
    m = re.match(r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", v)
    if m:
        return int(m.group(1)), int(m.group(2)), int(m.group(3))
    return None


def _parse_table_rows(lines: list[str]) -> list[list[str]]:
    """解析 GFM 表格为二维列表，自动跳过对齐分隔行。"""
    rows: list[list[str]] = []
    for ln in lines:
        ln = ln.strip()
        if not ln:
            continue
        if _BLOCK_TABLE_SEP_RE.match(ln):
            continue
        cells = _split_gfm_cells(ln)
        if cells:
            rows.append(cells)
    return rows


def _split_gfm_cells(line: str) -> list[str]:
    """拆分 GFM 表格行，处理 \\| 转义与 $...$ 内管道字符。"""
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    cells: list[str] = []
    depth = 0
    buf: list[str] = []
    i = 0
    while i < len(inner):
        ch = inner[i]
        if ch == "\\" and i + 1 < len(inner) and inner[i + 1] == "|":
            buf.append("|")
            i += 2
        elif ch == "$":
            depth ^= 1
            buf.append(ch)
            i += 1
        elif ch == "|" and depth == 0:
            cells.append("".join(buf).strip())
            buf = []
            i += 1
        else:
            buf.append(ch)
            i += 1
    cells.append("".join(buf).strip())
    return cells


# ──────────────────────────────────────────────
# 图片获取（本地优先，失败再下载）
# ──────────────────────────────────────────────

def _resolve_image(src: str, base_html: Path) -> Path | None:
    """解析 <img src> 为可读取的本地图片路径。

    本地图片模式下 src 已是本地相对路径；否则尝试按 base_html 的目录
    解析，再用网络下载到临时目录（仅导出时联网一次）。
    """
    if not src:
        return None
    # 已是本地路径
    cand = (base_html.parent / src) if not src.startswith(("http://", "https://")) else None
    if cand and cand.exists():
        return cand
    if src.startswith(("http://", "https://")):
        try:
            import tempfile
            import uuid
            suffix = Path(src.split("?")[0]).suffix or ".png"
            tmp = Path(tempfile.gettempdir()) / f"zhimg_{uuid.uuid4().hex}{suffix}"
            urllib.request.urlretrieve(src, tmp)  # noqa: S310
            if tmp.exists():
                return tmp
        except Exception as e:
            logger.warning(f"图片下载失败，跳过: {src} ({e})")
    return None


# ──────────────────────────────────────────────
# HTML → DOCX（python-docx，支持合并单元格 / 颜色 / 内嵌图）
# ──────────────────────────────────────────────

def _replace_math_with_alttext(node) -> None:
    """将节点内所有 <math> 替换为含其 alttext 的 span，
    避免后续 get_text() 同时输出 MathML 文本与 annotation 源码。
    """
    for m in node.find_all("math"):
        tex, _ = _extract_math_tex(m)
        replacement_text = tex or m.get_text(" ", strip=True) or ""
        span = BeautifulSoup("", "html.parser").new_tag("span")
        span.string = replacement_text
        m.replace_with(span)


def _extract_math_tex(math_node) -> tuple[str | None, bool]:
    """从 <math> 节点提取 LaTeX 源码与是否为 display 模式。"""
    tex = None
    ann = math_node.find("annotation", attrs={"encoding": "application/x-tex"})
    if ann and ann.get_text(strip=True):
        tex = ann.get_text(strip=True)
    if tex is None:
        alt = math_node.get("alttext")
        if alt:
            tex = alt
    disp = (math_node.get("display") or "").lower() == "block"
    # 外层 span.ltx_Math 可能带 ltx_display 标记
    parent = math_node.find_parent(class_="ltx_Math")
    if parent is not None and "ltx_display" in (parent.get("class") or []):
        disp = True
    return tex, disp


def _docx_inline_runs(p, node, base_html: Path) -> None:
    """把 HTML 节点（p / span / strong / em / code / a / math）的可见文本写到 docx 段落，
    逐节点保留粗体/斜体/颜色/下划线等内联样式，数学公式渲染为内嵌图片。
    """
    from docx.shared import Pt, RGBColor

    for child in node.children:
        if isinstance(child, str):
            txt = child.strip()
            if txt:
                run = p.add_run(txt)
                _apply_inline_style(run, child, RGBColor)
            continue
        name = child.name
        if name in ("br",):
            p.add_run("\n")
            continue
        if name in ("img",):
            img_path = _resolve_image(child.get("src", ""), base_html)
            if img_path:
                try:
                    from docx.shared import Cm
                    p.add_run().add_picture(str(img_path), width=Cm(14))
                except Exception:
                    p.add_run(f"[图: {child.get('src', '')}]")
            else:
                p.add_run(f"[图: {child.get('src', '')}]")
            continue
        # 数学公式：渲染为图片内嵌
        if name == "math" or (
            isinstance(cls := child.get("class"), list) and "ltx_Math" in cls
        ):
            math_el = child if name == "math" else child.find("math")
            if math_el is not None:
                tex, disp = _extract_math_tex(math_el)
                if tex:
                    from paper_tools.core.math_render import render_math_to_png

                    cache = Path(base_html).parent / "math"
                    png = render_math_to_png(tex, display=disp, out_dir=cache)
                    if png:
                        try:
                            from PIL import Image as _PILImage
                            from docx.shared import Cm

                            with _PILImage.open(png) as im:
                                iw, ih = im.size
                            # display 公式单独成行居中；inline 公式按图片宽高比
                            if disp:
                                para = p._parent.add_paragraph()
                                para.alignment = _docx_center()
                                para.paragraph_format.space_before = Cm(0.3)
                                para.paragraph_format.space_after = Cm(0.3)
                                run = para.add_run()
                                # display 公式：宽度限制 14cm，按原宽高比
                                if ih > 0 and iw > 0:
                                    target_w_cm = min(14.0, 14.0)
                                    target_h_cm = target_w_cm * (ih / iw)
                                    run.add_picture(str(png), width=Cm(target_w_cm))
                                else:
                                    run.add_picture(str(png), width=Cm(12))
                            else:
                                # inline 公式：以文字基线为目标高度(Cm(0.55)~15.6pt)，
                                # 按图片宽高比计算宽度
                                run = p.add_run()
                                run.font.raise_ = None  # 防止被 superscript 误判
                                if ih > 0 and iw > 0:
                                    target_h_cm = 0.55
                                    run.add_picture(str(png), height=Cm(target_h_cm))
                                else:
                                    run.add_picture(str(png), height=Cm(0.55))
                        except Exception:
                            p.add_run(f"[{tex}]")
                    else:
                        p.add_run(f"[{tex}]")
                continue
            # 无 tex 则跳过
            continue
        # 递归处理容器（strong/em/span/a/...）
        # 先写容器自身文本，再递归子节点
        direct = child.find(text=True, recursive=False)
        if direct and direct.strip():
            run = p.add_run(direct.strip())
            _apply_inline_style(run, child, RGBColor)
        for sub in child.children:
            if isinstance(sub, str):
                continue
            _docx_inline_runs(p, sub, base_html)


def _docx_center():
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    return WD_ALIGN_PARAGRAPH.CENTER



def _apply_inline_style(run, node, RGBColor) -> None:
    """根据 HTML 节点的 style / 标签给 run 设置粗体/斜体/颜色/字体。"""
    name = getattr(node, "name", None)
    if name in ("strong", "b"):
        run.bold = True
    if name in ("em", "i"):
        run.italic = True
    if name == "code":
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    if name == "a":
        run.font.color.rgb = RGBColor(0, 0, 200)
        run.underline = True
    style = node.get("style", "") if hasattr(node, "get") else ""
    if style:
        if re.search(r"font-weight\s*:\s*(bold|[6-9]00)", style, re.I):
            run.bold = True
        if re.search(r"font-style\s*:\s*italic", style, re.I):
            run.italic = True
        cm = re.search(r"color\s*:\s*([^;]+)", style, re.I)
        if cm:
            col = _parse_color(cm.group(1))
            if col:
                run.font.color.rgb = RGBColor(*col)


def _docx_add_table(doc, table_node, base_html: Path) -> None:
    """把 HTML <table>（支持 colspan/rowspan 合并）渲染为 docx 表格。"""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    rows = table_node.find_all("tr")
    if not rows:
        return
    # 计算逻辑网格尺寸
    n_cols = 0
    grid: list[list[dict]] = []
    for tr in rows:
        cells = [c for c in tr.find_all(["td", "th"])]
        logical_row: list[dict] = []
        col_cursor = 0
        for c in cells:
            # 跳过被上方合并占用的列
            while any(cell.get("_occupied", False) for cell in
                      [grid[r][col_cursor] for r in range(len(grid)) if col_cursor < len(grid[r])]):
                col_cursor += 1
            rs = int(c.get("rowspan", 1) or 1)
            cs = int(c.get("colspan", 1) or 1)
            # 标记占用
            for rr in range(rs):
                for cc in range(cs):
                    target_row = len(grid) + rr
                    while len(grid) <= target_row:
                        grid.append([])
                    while len(grid[target_row]) <= col_cursor + cc:
                        grid[target_row].append({})
                    grid[target_row][col_cursor + cc]["_occupied"] = True
            logical_row.append({
                "node": c, "rs": rs, "cs": cs,
                "col": col_cursor, "is_th": c.name == "th",
            })
            col_cursor += cs
        n_cols = max(n_cols, col_cursor)

    # 用稳定方式构建逻辑网格（含合并占位）
    logical = _compute_table_grid(rows)
    n_cols = max((len(r) for r in logical), default=0)
    n_rows = len(logical)
    if n_rows == 0 or n_cols == 0:
        return

    table = doc.add_table(rows=n_rows, cols=n_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(logical):
        for ci, cell in enumerate(row):
            # 跳过空格与被合并占用的占位格（合并由左上角单元格统一处理）
            if cell is None or cell.get("_span"):
                continue
            docx_cell = table.cell(ri, ci)
            c = cell["node"]
            # 合并单元格
            if cell["rs"] > 1 or cell["cs"] > 1:
                merge_end_row = min(ri + cell["rs"] - 1, n_rows - 1)
                merge_end_col = min(ci + cell["cs"] - 1, n_cols - 1)
                if merge_end_row > ri or merge_end_col > ci:
                    try:
                        docx_cell.merge(table.cell(merge_end_row, merge_end_col))
                    except Exception:
                        pass
            docx_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            _docx_inline_runs(docx_cell.paragraphs[0], c, base_html)
            for run in docx_cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                if cell["is_th"]:
                    run.bold = True


def _compute_table_grid(rows) -> list[list[dict | None]]:
    """计算表格的逻辑网格（含合并占位），返回 logical[row][col]。"""
    n = len(rows)
    # 先求列数
    max_cols = 0
    raw: list[list[dict]] = []
    for tr in rows:
        cells = [c for c in tr.find_all(["td", "th"])]
        logical_row: list[dict] = []
        col = 0
        for c in cells:
            # 跳过被上方合并占用的列
            while _cell_occupied(raw, len(raw), col):
                col += 1
            rs = int(c.get("rowspan", 1) or 1)
            cs = int(c.get("colspan", 1) or 1)
            logical_row.append({"node": c, "rs": rs, "cs": cs, "col": col,
                                "is_th": c.name == "th"})
            col += cs
        raw.append(logical_row)
        max_cols = max(max_cols, col)

    # 构建完整网格并填充合并
    grid: list[list[dict | None]] = [[None] * max_cols for _ in range(n)]
    for ri, logical_row in enumerate(raw):
        for cell in logical_row:
            c0 = cell["col"]
            for rr in range(cell["rs"]):
                for cc in range(cell["cs"]):
                    if ri + rr < n and c0 + cc < max_cols:
                        if rr == 0 and cc == 0:
                            grid[ri + rr][c0 + cc] = cell
                        else:
                            grid[ri + rr][c0 + cc] = {"_span": True}
    return grid


def _cell_occupied(raw: list[list[dict]], row_idx: int, col: int) -> bool:
    """判断 (row_idx, col) 是否已被上方单元格的 rowspan 占用。"""
    for r in range(row_idx):
        if r < len(raw):
            for cell in raw[r]:
                if cell and cell["col"] == col:
                    # 该 cell 的 rowspan 覆盖到当前行？
                    if (cell.get("rs", 1) or 1) > (row_idx - r):
                        return True
    return False


def _add_docx_heading(doc, text: str, level: int, base_html: Path) -> None:
    from docx.shared import Pt
    if level < 1:
        level = 1
    if level > 6:
        level = 6
    p = doc.add_heading("", level=level)
    # 把 text（纯文本）直接写，按标题层级
    p.add_run(text)
    sizes = {1: Pt(22), 2: Pt(18), 3: Pt(15), 4: Pt(13), 5: Pt(12), 6: Pt(11)}
    for run in p.runs:
        run.font.size = sizes.get(level, Pt(12))


def html_to_docx(html_path: Path, docx_path: Path | None = None) -> Path:
    """将保留原 HTML 结构的 `.zh.html` 导出为 Word (.docx)。

    支持：表格合并单元格、内联文字颜色、图片内嵌、公式（以文本呈现）。

    Args:
        html_path: `.zh.html` 文件（或降级使用的 `.zh.md`）。
        docx_path: 输出 .docx 路径；默认同目录、同 stem 的 .docx。

    Returns:
        输出的 .docx 文件路径。
    """
    html_path = Path(html_path)
    if docx_path is None:
        docx_path = html_path.with_suffix(".docx")

    if html_path.suffix.lower() == ".md":
        return _md_to_docx_fallback(html_path, docx_path)

    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    base_html = html_path
    _docx_equation_cache_dir["dir"] = (Path(html_path).parent / "math")
    _docx_equation_cache_dir["dir"] = (Path(html_path).parent / "math")

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rpr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), "宋体")

    # 论文正文容器（ltx_page_content）；找不到则全文档
    content = soup.find(class_="ltx_page_content") or soup.find("article") or soup.body
    if content is None:
        content = soup

    # 顺序渲染正文（保留结构）
    _render_docx_children(doc, content, base_html)

    try:
        doc.save(str(docx_path))
    except Exception:
        logger.exception("DOCX 导出失败")
        raise

    logger.info(f"DOCX 已导出: {docx_path}")
    return docx_path


def _render_docx_children(doc, parent, base_html: Path) -> None:
    """递归渲染 HTML 节点为 docx 内容（按正文顺序）。"""
    title_done = False
    for el in parent.children:
        if isinstance(el, str):
            continue
        if el.name is None:
            continue
        title_done = _render_docx_node(doc, el, base_html, title_done)
        if el.name in ("section", "div", "article", "blockquote"):
            # 递归内部（但 figure/table 已自身处理，避免重复）
            if el.name not in ("table", "figure"):
                _render_docx_children(doc, el, base_html)


def _render_docx_node(doc, el, base_html: Path, title_done: bool) -> bool:
    """渲染单个 HTML 节点；返回是否已处理标题（h1）。"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    name = el.name
    cls = " ".join(el.get("class", [])) if el.get("class") else ""

    # 跳过无关区域
    if any(s in cls for s in ("ltx_page_navbar", "ltx_bibliography",
                              "ltx_page_footer", "ltx_errors", "ltx_page_logo")):
        return title_done

    if name == "table" or "ltx_table" in cls:
        _docx_add_table(doc, el if name == "table" else el.find("table") or el, base_html)
        return title_done

    if name == "figure" or "ltx_figure" in cls:
        _render_docx_figure(doc, el, base_html)
        return title_done

    if "ltx_equation" in cls or "ltx_equationgroup" in cls:
        _render_docx_equation(doc, el)
        return title_done

    if "ltx_algorithm" in cls or "ltx_listing" in cls or name in ("pre", "code"):
        _render_docx_listing(doc, el)
        return title_done

    if name in ("h1", "h2", "h3", "h4", "h5", "h6") or "ltx_title" in cls:
        level = int(name[1]) if name.startswith("h") and name[1:].isdigit() else 2
        if "ltx_title_document" in cls:
            level = 1
        elif "ltx_title_section" in cls:
            level = 2
        elif "ltx_title_subsection" in cls:
            level = 3
        elif "ltx_title_subsubsection" in cls:
            level = 4
        elif "ltx_title_paragraph" in cls:
            level = 4
        elif "ltx_title_abstract" in cls:
            level = 2
        txt = el.get_text(" ", strip=True)
        if txt:
            _add_docx_heading(doc, txt, level, base_html)
        return title_done or (level == 1)

    if name == "p" or "ltx_p" in cls:
        txt = el.get_text(" ", strip=True)
        if not txt:
            return title_done
        p = doc.add_paragraph()
        _docx_inline_runs(p, el, base_html)
        return title_done

    if name in ("ul", "ol") or "ltx_itemize" in cls or "ltx_enumerate" in cls:
        for li in el.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            _docx_inline_runs(p, li, base_html)
        return title_done

    if name == "blockquote":
        for sub in el.children:
            if isinstance(sub, str):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            _docx_inline_runs(p, sub, base_html)
        return title_done

    # 其他容器：递归
    if name in ("div", "section", "article", "span"):
        return title_done

    return title_done


def _render_docx_figure(doc, fig, base_html: Path) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    img = fig.find("img")
    if img:
        src = img.get("src", "")
        img_path = _resolve_image(src, base_html)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_path:
            try:
                p.add_run().add_picture(str(img_path), width=Cm(14))
            except Exception:
                p.add_run(f"[图: {src}]")
        else:
            p.add_run(f"[图: {src}]")
    cap = fig.find(class_="ltx_caption") or fig.find("figcaption")
    if cap:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_inline_runs(cp, cap, base_html)
        for run in cp.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.italic = True


def _render_docx_equation(doc, el) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    math_el = el.find("math")
    tex = None
    if math_el is not None:
        tex, _ = _extract_math_tex(math_el)
    if tex is None:
        tex = el.get_text("\n", strip=True)
    if not tex:
        return
    from paper_tools.core.math_render import render_math_to_png

    cache = Path(_docx_equation_cache_dir.get("dir")) if _docx_equation_cache_dir else None
    png = render_math_to_png(tex, display=True, out_dir=cache)
    if png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(png), width=Cm(14))
        except Exception:
            run = p.add_run(tex)
            run.italic = True
            run.font.size = Pt(11)
    else:
        for part in tex.split("\n"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(90, 90, 90)
            run.italic = True


# DOCX 公式渲染使用的缓存目录（由 html_to_docx 设置）
_docx_equation_cache_dir: dict = {}


def _render_docx_listing(doc, el) -> None:
    from docx.shared import Pt

    code = el.get_text("\n", strip=True)
    if not code:
        return
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # 等宽中文字体回退
    rpr = run._element.get_or_add_rPr()
    from docx.oxml.ns import qn
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rpr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), "宋体")


def _md_to_docx_fallback(md_path: Path, docx_path: Path) -> Path:
    """基于 Markdown 的降级 DOCX 导出（无 HTML 时使用）。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.split("\n")
    n = len(lines)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, m, Cm(2.5))
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rpr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), "宋体")

    def _apply_inline(p, text):
        parts = re.split(r"(\*\*.*?\*\*|`[^`]+`|\$[^$\n]+\$|\[.*?\]\(.*?\))", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2]); run.bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1]); run.font.name = "Consolas"; run.font.size = Pt(10)
            elif part.startswith("$") and part.endswith("$"):
                run = p.add_run(part); run.italic = True; run.font.color.rgb = RGBColor(100,100,100)
            elif part.startswith("[") and "](" in part:
                m = _LINK_RE.match(part)
                if m:
                    run = p.add_run(m.group(1)); run.font.color.rgb = RGBColor(0,0,200); run.underline = True
            else:
                p.add_run(part)

    def _collect_table(start):
        buf = []
        j = start
        while j < n and _BLOCK_TABLE_ROW_RE.match(lines[j]):
            buf.append(lines[j]); j += 1
        return _parse_table_rows(buf), j

    h1_done = False
    i = 0
    while i < n:
        line = lines[i].rstrip(); i += 1
        if not line.strip():
            continue
        if _BLOCK_HR_RE.match(line.strip()):
            doc.add_paragraph("─" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        hm = _BLOCK_HEADING_RE.match(line)
        if hm:
            level = len(hm.group(1))
            txt = _strip_inline(hm.group(2))
            if level == 1 and not h1_done:
                _add_docx_heading(doc, txt, 1, md_path); h1_done = True
            elif level != 1:
                _add_docx_heading(doc, txt, level, md_path)
            continue
        if _BLOCK_BLOCKQUOTE_RE.match(line):
            i -= 1
            buf = []
            while i < n and _BLOCK_BLOCKQUOTE_RE.match(lines[i]):
                buf.append(_BLOCK_BLOCKQUOTE_RE.match(lines[i]).group(1)); i += 1
            for ln in buf:
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.0)
                run = p.add_run(ln); run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(120,120,120); run.italic = True
            continue
        if _BLOCK_TABLE_ROW_RE.match(line):
            i -= 1
            rows, ni = _collect_table(i); i = ni
            if rows:
                ncols = max(len(r) for r in rows)
                padded = [r + [""] * (ncols - len(r)) for r in rows]
                table = doc.add_table(rows=len(padded), cols=ncols, style="Table Grid")
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row_data in enumerate(padded):
                    for ci, cell_text in enumerate(row_data):
                        p = table.cell(ri, ci).paragraphs[0]
                        run = p.add_run(_strip_inline(cell_text)); run.font.size = Pt(10)
                        if ri == 0:
                            run.bold = True
            continue
        if line.startswith("$$") and line.endswith("$$"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line); run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(90,90,90); run.italic = True
            continue
        p = doc.add_paragraph()
        _apply_inline(p, line)

    doc.save(str(docx_path))
    logger.info(f"DOCX 已导出（基于 markdown 降级）: {docx_path}")
    return docx_path


# ──────────────────────────────────────────────
# HTML → PDF（fpdf2，支持合并单元格 / 颜色 / 内嵌图）
# ──────────────────────────────────────────────

_CN_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\simsun.ttc", r"C:\Windows\Fonts\simsunb.ttf",
    r"C:\Windows\Fonts\SimsunExtG.ttf", r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\simfang.ttf", r"C:\Windows\Fonts\STFANGSO.TTF",
    r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\Deng.ttf",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/PingFang.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
]

_MONO_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\consola.ttf", r"C:\Windows\Fonts\cour.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
    "/usr/share/fonts/TTF/DejaVuSansMono.ttf",
]


def _find_font(candidates: list[str]) -> str | None:
    for path in candidates:
        if Path(path).exists():
            return path
    return None


def html_to_pdf(html_path: Path, pdf_path: Path | None = None) -> Path:
    """将保留原 HTML 结构的 `.zh.html` 导出为 PDF。

    支持：表格合并、内联文字颜色、图片内嵌。

    Args:
        html_path: `.zh.html` 文件（或降级使用的 `.zh.md`）。
        pdf_path: 输出 PDF 路径；默认同目录、同 stem 的 .pdf。

    Returns:
        输出的 PDF 文件路径。
    """
    html_path = Path(html_path)
    if pdf_path is None:
        pdf_path = html_path.with_suffix(".pdf")

    if html_path.suffix.lower() == ".md":
        return _md_to_pdf_fallback(html_path, pdf_path)

    from fpdf import FPDF

    cn_font = _find_font(_CN_FONT_CANDIDATES)
    if cn_font is None:
        raise FileNotFoundError(
            "无法找到中文字体文件。请确保系统已安装宋体/黑体/微软雅黑之一。\n"
            f"搜索路径: {_CN_FONT_CANDIDATES}"
        )
    heading_font = cn_font
    for path in _CN_FONT_CANDIDATES:
        if "simhei" in path.lower() or "hei" in path.lower():
            if Path(path).exists():
                heading_font = path
                break
    italic_font = cn_font
    for path in _CN_FONT_CANDIDATES:
        if "fang" in path.lower():
            if Path(path).exists():
                italic_font = path
                break
    mono_font = _find_font(_MONO_FONT_CANDIDATES)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.add_font("CNBody", "", cn_font, uni=True)
    pdf.add_font("CNBody", "B", heading_font, uni=True)
    pdf.add_font("CNBody", "I", italic_font, uni=True)
    pdf.add_font("CNBody", "BI", heading_font, uni=True)
    if mono_font:
        pdf.add_font("CNMono", "", mono_font, uni=True)

    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    base_html = html_path
    _docx_equation_cache_dir["dir"] = (Path(html_path).parent / "math")
    content = soup.find(class_="ltx_page_content") or soup.find("article") or soup.body or soup

    page_width = pdf.w - pdf.l_margin - pdf.r_margin

    def _rendernode(node, in_list=False):
        for child in node.children:
            if isinstance(child, str):
                continue
            if child.name is None:
                continue
            cls = " ".join(child.get("class", [])) if child.get("class") else ""
            if any(s in cls for s in ("ltx_page_navbar", "ltx_bibliography",
                                      "ltx_page_footer", "ltx_errors", "ltx_page_logo")):
                continue
            if child.name == "table" or "ltx_table" in cls:
                _pdf_table(pdf, child if child.name == "table" else child.find("table") or child,
                           page_width, base_html)
                continue
            if child.name == "figure" or "ltx_figure" in cls:
                _pdf_figure(pdf, child, page_width, base_html)
                continue
            if "ltx_equation" in cls or "ltx_equationgroup" in cls:
                # 提取内部 <math> 的 LaTeX 源码渲染为图片（display）
                math_el = child.find("math")
                tex = None
                if math_el is not None:
                    tex, _ = _extract_math_tex(math_el)
                if tex is None:
                    tex = child.get_text("\n", strip=True)
                if tex:
                    from paper_tools.core.math_render import render_math_to_png

                    cache = Path(base_html).parent / "math"
                    png = render_math_to_png(tex, display=True, out_dir=cache)
                    if png:
                        _pdf_insert_math(pdf, png, display=True,
                                         page_width=page_width, indent=0)
                    else:
                        pdf.ln(2)
                        pdf.set_font("CNBody", "I", 10)
                        pdf.set_text_color(90, 90, 90)
                        for part in tex.split("\n"):
                            pdf.multi_cell(page_width, 7, part, align="C",
                                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                        pdf.set_text_color(0, 0, 0)
                        pdf.ln(2)
                continue
            if "ltx_algorithm" in cls or "ltx_listing" in cls or child.name in ("pre", "code"):
                code = child.get_text("\n", strip=True)
                if code:
                    pdf.set_font("CNMono", "", 9)
                    pdf.set_text_color(60, 60, 60)
                    pdf.multi_cell(page_width, 5, code,
                                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.set_text_color(0, 0, 0)
                    pdf.ln(2)
                continue
            if child.name in ("h1", "h2", "h3", "h4", "h5", "h6") or "ltx_title" in cls:
                level = 2
                if "ltx_title_document" in cls:
                    level = 1
                elif "ltx_title_section" in cls:
                    level = 2
                elif "ltx_title_subsection" in cls:
                    level = 3
                elif "ltx_title_subsubsection" in cls or "ltx_title_paragraph" in cls:
                    level = 4
                elif "ltx_title_abstract" in cls:
                    level = 2
                sizes = {1: 18, 2: 15, 3: 13, 4: 12}
                sz = sizes.get(level, 12)
                txt = child.get_text(" ", strip=True)
                if txt:
                    pdf.ln(3)
                    pdf.set_font("CNBody", "B", sz)
                    pdf.multi_cell(page_width, sz * 0.7, txt,
                                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.ln(2)
                continue
            if child.name == "p" or "ltx_p" in cls:
                txt = child.get_text(" ", strip=True)
                if txt:
                    _pdf_rich_line(pdf, child, page_width, base_html)
                continue
            if child.name in ("ul", "ol") or "ltx_itemize" in cls or "ltx_enumerate" in cls:
                for li in child.find_all("li", recursive=False):
                    _pdf_rich_line(pdf, li, page_width, base_html, indent=6)
                continue
            if child.name == "blockquote":
                for sub in child.children:
                    if isinstance(sub, str):
                        continue
                    _pdf_rich_line(pdf, sub, page_width, base_html, indent=8,
                                  color=(120, 120, 120), italic=True)
                continue
            # 容器：递归
            if child.name in ("div", "section", "article", "span"):
                _rendernode(child, in_list)

    _rendernode(content)

    pdf.output(str(pdf_path))
    logger.info(f"PDF 已导出: {pdf_path}")
    return pdf_path


def _pdf_rich_line(pdf, node, page_width, base_html: Path, indent=0,
                   color=None, italic=False) -> None:
    """渲染一个带内联样式的段落文本到 PDF（保留颜色 / 粗体 / 斜体）。

    使用 `write()` 流式输出，各片段可带不同字体/颜色且能正确自动换行与
    分页；避免逐片段 `multi_cell` 导致的异常分页（每片段独占一页）。
    """
    segs = _collect_inline_segments(node)
    if not segs:
        return
    pdf.ln(1)
    if indent:
        pdf.set_x(pdf.l_margin + indent)
    for seg in segs:
        text, bold, it, col, kind = seg
        if kind == "br":
            pdf.ln(6)
            if indent:
                pdf.set_x(pdf.l_margin + indent)
            continue
        if kind == "img":
            img_path = _resolve_image(text, base_html)
            if img_path:
                try:
                    pdf.image(str(img_path), w=page_width - indent)
                except Exception:
                    pdf.write(6, f"[图: {text}]")
            else:
                pdf.write(6, f"[图: {text}]")
            continue
        if kind == "math":
            from paper_tools.core.math_render import render_math_to_png

            cache = Path(base_html).parent / "math"
            png = render_math_to_png(text, display=bold, out_dir=cache)
            if png:
                _pdf_insert_math(pdf, png, display=bold, page_width=page_width,
                                 indent=indent)
                continue
            # 渲染失败则退化为源码文本
            pdf.set_font("CNBody", "I", 11)
            pdf.set_text_color(90, 90, 90)
            pdf.write(6, f"[{text}]")
            pdf.set_text_color(0, 0, 0)
            continue
        # 普通文本
        style = ""
        if bold:
            style += "B"
        if it or italic:
            style += "I"
        pdf.set_font("CNBody", style, 11)
        if col:
            pdf.set_text_color(*col)
        elif color:
            pdf.set_text_color(*color)
        else:
            pdf.set_text_color(0, 0, 0)
        pdf.write(6, text)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(7)  # 段落结束换行


def _pdf_insert_math(pdf, png: Path, display: bool, page_width: float,
                     indent: int = 0) -> None:
    """在 PDF 流中插入公式图片：inline 随文，display 单独成行居中。"""
    from PIL import Image

    with Image.open(png) as im:
        iw, ih = im.size
    if iw <= 0 or ih <= 0:
        return
    # inline 基准：以 11pt 字号(~6mm 行高)对应的像素高度作为缩放依据
    line_mm = 6.0
    if display:
        h_mm = 12.0
        w_mm = h_mm * (iw / ih)
        if w_mm > page_width - indent:
            w_mm = page_width - indent
            h_mm = w_mm * (ih / iw)
        # 段落内 display：用稍小的前置间距避免与上方文字重叠
        y_cursor_before = pdf.get_y()
        # 若公式总高超过当前可用区域，先分页
        if y_cursor_before + h_mm + 4 > pdf.h - pdf.b_margin:
            pdf.add_page()
        pdf.ln(2)
        y = pdf.get_y()
        x = pdf.l_margin + indent + max(0, (page_width - indent - w_mm) / 2)
        pdf.image(str(png), x=x, y=y, w=w_mm)
        # 显式把 y 推到公式底部，确保后续文字不会覆盖
        pdf.set_y(y + h_mm)
        pdf.ln(2)
    else:
        # inline：目标高度 = 行高，再按原图宽高比算宽度
        ratio = ih / iw
        # 想要公式与文字行同高（约 5mm），按比例得出宽度
        target_h_mm = 5.0
        h_mm = target_h_mm
        w_mm = h_mm / ratio if ratio > 0 else target_h_mm
        # 上限：行宽的 70%，避免单行只塞一个公式
        max_w = (page_width - indent) * 0.7
        if w_mm > max_w:
            w_mm = max_w
            h_mm = w_mm * ratio
        # 若当前行放不下则换行
        if pdf.get_x() + w_mm > pdf.l_margin + page_width:
            pdf.ln(line_mm)
            pdf.set_x(pdf.l_margin + indent)
        # 垂直基线偏移：把图片底部对齐到文字基线，顶部稍高
        y = pdf.get_y() + (line_mm - h_mm) * 0.4
        if y < pdf.t_margin:
            y = pdf.t_margin
        pdf.image(str(png), x=pdf.get_x(), y=y, w=w_mm)
        pdf.set_x(pdf.get_x() + w_mm)
        # 如果公式比正文行高更高，把行基线向下推
        if h_mm + 1 > line_mm:
            pdf.set_y(pdf.get_y() + (h_mm - line_mm) + 1)


def _collect_inline_segments(node) -> list[tuple]:
    """递归收集 HTML 节点的可见文本片段。

    每个片段为 5 元组 (text, bold, italic, color, kind)：
      - kind="text"：普通文本，text 为字符串
      - kind="math"：数学公式，text 为 LaTeX 源码，bold 字段复用为 display 布尔
      - kind="img" ：图片占位，text 为 src
    此外 "br" 以 ("\n",...) 表示。
    """
    segs: list[tuple] = []
    for child in node.children:
        if isinstance(child, str):
            t = child.strip()
            if t:
                col = _segment_color(node)
                segs.append((t, _segment_bold(node), _segment_italic(node), col, "text"))
            continue
        name = child.name
        if name == "br":
            segs.append(("\n", False, False, None, "br"))
            continue
        if name == "img":
            segs.append((child.get("src", ""), False, True, (120, 120, 120), "img"))
            continue
        # 数学公式：提取 LaTeX 源码，渲染为图片
        if name == "math" or (
            isinstance(cls := child.get("class"), list) and "ltx_Math" in cls
        ):
            math_el = child if name == "math" else child.find("math")
            if math_el is not None:
                tex, disp = _extract_math_tex(math_el)
                if tex:
                    segs.append((tex, disp, False, None, "math"))
            continue
        # 容器：递归
        segs.extend(_collect_inline_segments(child))
    return segs


def _segment_bold(node) -> bool:
    if node.name in ("strong", "b"):
        return True
    style = node.get("style", "") if hasattr(node, "get") else ""
    return bool(re.search(r"font-weight\s*:\s*(bold|[6-9]00)", style, re.I))


def _segment_italic(node) -> bool:
    if node.name in ("em", "i"):
        return True
    style = node.get("style", "") if hasattr(node, "get") else ""
    return bool(re.search(r"font-style\s*:\s*italic", style, re.I))


def _segment_color(node) -> tuple | None:
    style = node.get("style", "") if hasattr(node, "get") else ""
    m = re.search(r"color\s*:\s*([^;]+)", style, re.I)
    if m:
        return _parse_color(m.group(1))
    return None


def _pdf_figure(pdf, fig, page_width, base_html: Path) -> None:
    img = fig.find("img")
    if img:
        src = img.get("src", "")
        img_path = _resolve_image(src, base_html)
        if img_path:
            try:
                from PIL import Image
                with Image.open(img_path) as im:
                    w, h = im.size
                max_w = page_width
                ratio = h / w if w else 1
                pdf_w = max_w
                pdf_h = pdf_w * ratio
                if pdf_h > 220:
                    pdf_h = 220
                    pdf_w = pdf_h / ratio if ratio else max_w
                # 若当前剩余空间连图都放不下，先翻页
                if pdf.get_y() + pdf_h > pdf.h - pdf.b_margin:
                    pdf.add_page()
                x = pdf.l_margin + (page_width - pdf_w) / 2
                pdf.image(str(img_path), x=x, w=pdf_w)
                # 图片后让光标停止在图片底部，额外留 3mm 即可，不要再追加 pdf_h
                pdf.set_y(pdf.get_y() + 3)
                pdf.set_x(pdf.l_margin)
            except Exception:
                pdf.set_font("CNBody", "I", 9)
                pdf.set_text_color(120, 120, 120)
                pdf.multi_cell(page_width, 6, f"[图: {src}]",
                               new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.set_text_color(0, 0, 0)
        else:
            pdf.set_font("CNBody", "I", 9)
            pdf.set_text_color(120, 120, 120)
            pdf.multi_cell(page_width, 6, f"[图: {src}]",
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(0, 0, 0)
    cap = fig.find(class_="ltx_caption") or fig.find("figcaption")
    if cap:
        pdf.set_font("CNBody", "I", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(page_width, 6, cap.get_text(" ", strip=True), align="C",
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        # 图注后给整段后留 4mm 即可，避免大段空白
        pdf.ln(4)


def _pdf_table(pdf, table_node, page_width, base_html: Path) -> None:
    """把 HTML 表格（支持 colspan/rowspan）绘制到 PDF。

    使用 fpdf2 内置 table API：自动处理文字换行、单元格内边距、
    跨页分页与边框；colspan 通过 cell 的 colspan 参数实现，rowspan 通过
    记录"待续占位"在后续行用空 cell 占位实现。

    单元格内公式处理策略：
      - 若 cell 内**只有 1 个 <math>** ：用 fpdf2 cell 的 img 参数嵌入 PNG（自动缩放）
      - 若 cell 内**混合文本+math**（如 "minimize" 这种带文字）：渲染图片+纯文本退化为 alttext
    """
    logical = _compute_table_grid(table_node.find_all("tr"))
    if not logical:
        return
    n_cols = max((len(r) for r in logical), default=0)
    if n_cols == 0:
        return

    from paper_tools.core.math_render import render_math_to_png
    import copy as _copy
    import re as _re_table

    cache_dir = Path(base_html).parent / "math"

    def _math_to_unicode(tex: str) -> str:
        """把常见简单 LaTeX 公式转 unicode 字符——作为混合文字 cell 的退化方案。
        仅处理纯 ASCII 字母 + 上下标的常见情形（mathcal, mathbb, mathbf 等）。"""
        s = tex.strip()
        # \\mathcal{G} → 𝒢, \\mathcal{L} → ℒ, 等
        s = _re_table.sub(r"\\mathcal\{([A-Za-z])\}", _unicode_for, s)
        s = _re_table.sub(r"\\mathbb\{([A-Za-z])\}", _unicode_for, s)
        s = _re_table.sub(r"\\mathbf\{([A-Za-z])\}", lambda m: m.group(1).upper(), s)
        s = _re_table.sub(r"\\mathrm\{([A-Za-z])\}", lambda m: m.group(1), s)
        s = _re_table.sub(r"\\text\{([^{}]*)\}", r"\1", s)
        s = _re_table.sub(r"\\mathit\{([A-Za-z])\}", lambda m: m.group(1), s)
        s = _re_table.sub(r"\\mathfrak\{([A-Za-z])\}", lambda m: m.group(1), s)
        # 上下标简单化
        s = _re_table.sub(r"_\{([^{}]*)\}", r"\1", s)
        s = _re_table.sub(r"\^\{([^{}]*)\}", r"\1", s)
        s = _re_table.sub(r"\\\\quad|\\\\,|\\\\;|\\\\:|\\\\!", " ", s)
        s = s.replace("\\geq", ">=").replace("\\leq", "<=")
        s = s.replace("\\approx", "≈").replace("\\sim", "~")
        s = s.replace("\\times", "×").replace("\\cdot", "·")
        s = s.replace("\\leftarrow", "←").replace("\\rightarrow", "→")
        s = s.replace("\\infty", "∞")
        # 移除其他未识别指令
        s = _re_table.sub(r"\\[a-zA-Z]+\s*", "", s)
        s = _re_table.sub(r"[{}]", "", s)
        return s.strip() or tex

    _SCRIPT_UNICODE = {
        "G": "\U0001D4A2", "L": "\u2112", "P": "\U0001D4AB",
        "Q": "\U0001D4AC", "B": "\u212C", "F": "\u212D",
        "R": "\u211D", "C": "\u2102", "H": "\u210D",
        "I": "\u2110", "N": "\u2115", "Z": "\u2124",
    }

    def _unicode_for(m):
        ch = m.group(1).upper()
        return _SCRIPT_UNICODE.get(ch, ch)

    def _render_mixed_cell_png(cell_node, font_path: Path, target_w_px: int,
                                cache_dir: Path, dpi: int = 200) -> "Path | None":
        """把 mixed cell（含文字+多个 inline math）渲染为一张 PNG。

        关键改进：
          * target_w_px 是单元格所在列的实际像素宽度（而非整页宽）。
          * 文字按 target_w_px 自动换行，不再为了塞入一行而整体缩小。
          * 公式随所在行高度等比缩放。
          * 图片高度按内容自适应（多行累加）。
        返回 PNG 路径；若失败返回 None。
        """
        import hashlib as _h
        import re as _re_mc
        from PIL import Image as _PI, ImageDraw as _PD, ImageFont as _IF
        from paper_tools.core.math_render import render_math_to_png

        if target_w_px <= 0:
            target_w_px = 300

        # 先把所有 <math> 渲染为 PNG，并收集信息
        n = _copy.deepcopy(cell_node)
        math_imgs: list[tuple[int, Path, float]] = []
        for idx, mn in enumerate(n.find_all("math")):
            tex, disp = _extract_math_tex(mn)
            if not tex:
                mn.decompose()
                continue
            png_path = render_math_to_png(tex, display=False, out_dir=cache_dir)
            if not png_path:
                mn.decompose()
                continue
            try:
                imgm = _PI.open(png_path)
                imgm.load()
                ratio = imgm.size[1] / max(1, imgm.size[0])
                imgm.close()
            except Exception:
                ratio = 0.4
            mn.replace_with(f"⟦M{idx}⟧")
            math_imgs.append((idx, png_path, ratio))

        # 提取纯文本并清除 <img>（用 alt 替代）
        for img in list(n.find_all("img")):
            src = img.get("src", "")
            alt = img.get("alt", "")
            img_path = _resolve_image(src, base_html) if src else None
            if img_path and img_path.exists():
                try:
                    imgm = _PI.open(img_path)
                    imgm.load()
                    imgm.close()
                except Exception:
                    img_path = None
            if alt:
                img.replace_with(f"[{alt}]")
            elif img_path:
                img.replace_with("[图]")
            else:
                img.decompose()

        text = n.get_text(" ", strip=True)
        if not text:
            return None

        # 基准字号：11pt 在给定 dpi 下的像素尺寸
        base_font_px = max(7, int(round(11 * dpi / 72)))
        try:
            base_font = _IF.truetype(str(font_path), base_font_px)
        except Exception:
            base_font = _IF.load_default()

        # 把占位符还原成具体片段：文本或公式
        parts_text: list[str | None] = []
        parts_img: list[tuple[Path, float] | None] = []
        tokens = _re_mc.split(r"(⟦M\d+⟧)", text)
        for tok in tokens:
            if not tok:
                continue
            if tok.startswith("⟦M") and tok.endswith("⟧"):
                idx = int(tok[2:-1])
                if idx < len(math_imgs):
                    _, pp, ratio = math_imgs[idx]
                    parts_text.append(None)
                    parts_img.append((pp, ratio))
                continue
            parts_text.append(tok)
            parts_img.append(None)

        # 把每个片段拆成“原子”：文本按空格拆成 words，公式作为独立原子
        # 原子类型：("text", word) / ("img", pp, ratio)
        atoms: list[tuple[str, ...]] = []
        for txt, img_info in zip(parts_text, parts_img):
            if txt is not None:
                # 保留空格作为独立原子，便于换行时保留词间距
                for word in _re_mc.split(r"( +)", txt):
                    if not word:
                        continue
                    atoms.append(("text", word))
            elif img_info is not None:
                atoms.append(("img", img_info[0], img_info[1]))

        # 计算每个原子在基准字号下的尺寸
        atom_dims: list[tuple[str, int, int, tuple]] = []
        for atom in atoms:
            if atom[0] == "text":
                word = atom[1]
                w = int(base_font.getlength(word))
                bbox = base_font.getbbox(word)
                h = bbox[3] - bbox[1] if bbox else base_font_px
                atom_dims.append(("text", w, h, (word,)))
            else:
                _, pp, ratio = atom
                try:
                    imgm = _PI.open(pp)
                    imgm.load()
                    iw, ih = imgm.size
                    imgm.close()
                except Exception:
                    iw, ih = int(base_font_px * 2), base_font_px
                nh = int(round(base_font_px * 1.2))
                nw = max(1, int(round(nh / max(0.1, ratio))))
                atom_dims.append(("img", nw, nh, (pp, ratio)))

        if not atom_dims:
            return None

        padding_px = int(round(base_font_px * 0.5))
        usable_w = max(1, target_w_px - padding_px)

        # 检查是否单个原子就超出列宽（极窄列），需要整体缩小
        max_atom_w = max(w for _, w, _, _ in atom_dims)
        global_scale = 1.0
        if max_atom_w > usable_w:
            global_scale = usable_w / max_atom_w
        font_px = max(7, int(round(base_font_px * global_scale)))
        try:
            font = _IF.truetype(str(font_path), font_px)
        except Exception:
            font = base_font

        # 重新计算实际字号下的原子尺寸
        atoms2: list[tuple[str, int, int, tuple]] = []
        for typ, _, _, payload in atom_dims:
            if typ == "text":
                word = payload[0]
                w = int(font.getlength(word))
                bbox = font.getbbox(word)
                h = bbox[3] - bbox[1] if bbox else font_px
                atoms2.append(("text", w, h, (word,)))
            else:
                pp, ratio = payload
                try:
                    imgm = _PI.open(pp)
                    imgm.load()
                    iw, ih = imgm.size
                    imgm.close()
                except Exception:
                    iw, ih = int(font_px * 2), font_px
                nh = int(round(font_px * 1.2))
                nw = max(1, int(round(nh / max(0.1, ratio))))
                atoms2.append(("img", nw, nh, (pp, ratio)))

        usable_w = max(1, target_w_px - padding_px)

        # 自动换行：把原子排成多行，每行宽度不超过 usable_w
        lines: list[list[tuple[str, int, int, tuple]]] = [[]]
        cur_w = 0
        for atom in atoms2:
            aw = atom[1]
            # 行首空格不要；非行首保留一个空格宽度
            if atom[0] == "text" and atom[3][0].strip() == "":
                if not lines[-1]:
                    continue
                aw = min(aw, max(1, usable_w - cur_w))
            if cur_w + aw + (2 if lines[-1] else 0) > usable_w and lines[-1]:
                lines.append([atom])
                cur_w = aw
            else:
                lines[-1].append(atom)
                cur_w += aw + (2 if len(lines[-1]) > 1 else 0)

        # 去掉空行
        lines = [ln for ln in lines if ln]
        if not lines:
            return None

        # 计算每行高度（取行内最大高度）
        line_heights = []
        line_widths = []
        for ln in lines:
            h = max(h for _, _, h, _ in ln)
            w = sum(w for _, w, _, _ in ln) + max(0, len(ln) - 1) * 2
            line_heights.append(int(round(h * 1.2)))
            line_widths.append(w)

        line_spacing = max(2, int(round(font_px * 0.15)))
        png_w_px = max(line_widths) + padding_px
        png_h_px = sum(line_heights) + (len(lines) - 1) * line_spacing + padding_px
        img = _PI.new("RGBA", (png_w_px, png_h_px), (255, 255, 255, 0))
        draw = _PD.Draw(img)

        y = padding_px // 2
        for ln, lh in zip(lines, line_heights):
            x = padding_px // 2
            # 行内垂直居中
            for typ, w, h, payload in ln:
                if typ == "text":
                    word = payload[0]
                    bbox = font.getbbox(word)
                    text_h = bbox[3] - bbox[1] if bbox else font_px
                    yy = y + (lh - text_h) // 2 - (bbox[1] if bbox else 0)
                    draw.text((x, yy), word, fill=(0, 0, 0, 255), font=font)
                else:
                    pp, ratio = payload
                    try:
                        im = _PI.open(pp).convert("RGBA")
                        im = im.resize((w, h), _PI.Resampling.LANCZOS)
                        yy = y + (lh - h) // 2
                        img.paste(im, (x, yy), im)
                        im.close()
                    except Exception:
                        pass
                x += w + 2
            y += lh + line_spacing

        h = _h.md5(
            img.tobytes()
            + b"|" + str(target_w_px).encode()
            + b"|" + str(dpi).encode()
            + b"|" + str(font_px).encode()
        ).hexdigest()[:12]
        out = cache_dir / f"cell_mixed_{h}.png"
        if not out.exists():
            img.save(out, format="PNG", dpi=(dpi, dpi))
        return out

    def _trim_png_alpha(png_path, pad_px=1, target_h_px=None, target_ratio_w_over_h=None):
        """裁掉 PNG 透明外圈，让图片内容紧贴边框。
        可选地：
          - target_h_px: resize 到固定高度，宽按等比（保持原宽高比）；
          - target_ratio_w_over_h: 强制宽高比 W/H 到该值（让 cell 高窄而稳定）。
        写入 cache_dir/math_trim/<sha>.png，返回新路径。
        """
        import hashlib as _h
        from PIL import Image as _PI2
        p = str(png_path)
        if not p or not Path(p).exists():
            return None
        try:
            im = _PI2.open(p).convert("RGBA")
        except Exception:
            return None
        bbox = im.getbbox()
        if not bbox:
            im.close()
            return None
        l, t, r, b = bbox
        W, H = im.size
        l = max(0, l - pad_px); t = max(0, t - pad_px)
        r = min(W, r + pad_px); b = min(H, b + pad_px)
        try:
            crop = im.crop((l, t, r, b))
            tw, th = crop.size
            # 1) 等比缩放到固定高度
            if target_h_px and target_h_px > 0 and th > 0:
                scale = target_h_px / th
                nw = max(1, int(round(tw * scale)))
                nh = target_h_px
                crop = crop.resize((nw, nh), _PI2.LANCZOS)
                tw, th = nw, nh
            # 2) 强制宽高比（不放大，只调整空白画布，使所有图宽高比一致，
            #    这样被 img_fill_width 拉伸后 cell 高度统一）
            if target_ratio_w_over_h and target_ratio_w_over_h > 0 and th > 0:
                target_w = max(tw, int(round(th * target_ratio_w_over_h)))
                target_h = th
                # 若当前已宽，按目标比例放大宽
                if target_w > tw:
                    # 在右侧加透明 padding 使宽高比达标
                    canvas = _PI2.new("RGBA", (target_w, target_h), (255, 255, 255, 0))
                    canvas.paste(crop, (0, 0), crop)
                    crop = canvas
                    tw = target_w
            trim_dir = cache_dir / "math_trim"
            trim_dir.mkdir(parents=True, exist_ok=True)
            key = _h.md5(f"{p}|{l}|{t}|{r}|{b}|{target_h_px}|{target_ratio_w_over_h}".encode()).hexdigest()[:12]
            out = trim_dir / f"trim_{key}.png"
            crop.save(out, "PNG")
            im.close()
            return out
        except Exception:
            im.close()
            return Path(p)

    def _cell_pure_math(cell_node):
        """检测 cell 是否仅仅包含 1 个 <math> 节点（无其他文字/标签兄弟）。
        返回 (tex, disp, png_path) 或 None。

        判断方法：把 cell 下所有 math 节点抽掉后，看剩余非空文本/标签，
        若全为空则视为纯公式 cell（此时取该 math 的 PNG 渲染）。"""
        node = cell_node
        math_nodes = node.find_all("math")
        if len(math_nodes) != 1:
            return None
        # 去除所有 <math> 后，剩余应无可见内容（允许 <br> 或纯空白）
        clone = _copy.deepcopy(node)
        for m in clone.find_all("math"):
            m.decompose()
        clone.attrs = {}
        # 检查直接子元素：移除空白字符串和 <br>，若有剩余则非纯公式
        significant = []
        for c in clone.children:
            if isinstance(c, str):
                if c.strip():
                    significant.append(c.strip())
            else:
                if c.name == "br":
                    continue
                text = c.get_text(strip=True)
                if text:
                    significant.append(text)
        if significant:
            return None
        tex, disp = _extract_math_tex(math_nodes[0])
        if not tex:
            return None
        png = render_math_to_png(tex, display=disp, out_dir=cache_dir)
        if not png:
            return None
        # render_math_to_png 的 DPI 较高，直接嵌入会非常大；
        # 用 Pillow 显式裁剪透明边并缩放到固定显示高度（约 16px）。
        def _resize_math_for_cell(png_path: Path, target_h_px: int = 16) -> Path:
            from PIL import Image as _PI
            try:
                with _PI.open(png_path) as im:
                    im = im.convert("RGBA")
                    # 裁剪透明边
                    alpha = im.split()[-1]
                    bbox = alpha.getbbox()
                    if bbox:
                        im = im.crop(bbox)
                    # 等比缩放到目标高度
                    w, h = im.size
                    if h > target_h_px:
                        new_h = target_h_px
                        new_w = max(1, int(round(w * new_h / h)))
                        im = im.resize((new_w, new_h), _PI.Resampling.LANCZOS)
                    # 保存到新路径（避免覆盖原缓存）
                    import hashlib as _h2
                    key = _h2.md5(
                        (png_path.read_bytes()
                         if png_path.exists() else b"")
                        + b"|cellh=" + str(target_h_px).encode()
                    ).hexdigest()[:16]
                    out = cache_dir / f"cell_math_{key}.png"
                    if not out.exists():
                        im.save(out, format="PNG")
                    return out
            except Exception:
                return png_path

        resized = _resize_math_for_cell(png, target_h_px=16)
        return tex, disp, (resized or png)

    pdf.ln(2)

    def _parse_css_width(value: str) -> float | None:
        """解析 width 样式/属性；返回 mm（绝对单位）或 0-1 小数（百分比）。"""
        if not value:
            return None
        m = re.search(r"width:\s*([0-9.]+)\s*(px|pt|cm|mm|in|em|%|)", value, re.I)
        if not m:
            m = re.match(r"^\s*([0-9.]+)\s*(px|pt|cm|mm|in|em|%|)\s*$", value, re.I)
        if not m:
            return None
        val = float(m.group(1))
        unit = (m.group(2) or "").lower()
        if unit == "%":
            return val / 100.0
        if unit == "px" or unit == "":
            return val * 25.4 / 96.0
        if unit == "pt":
            return val * 25.4 / 72.0
        if unit == "cm":
            return val * 10.0
        if unit == "mm":
            return val
        if unit == "in":
            return val * 25.4
        if unit == "em":
            return val * 4.0
        return val * 25.4 / 96.0

    def _table_col_widths(table_node, total_mm: float, n_cols: int) -> list[float]:
        cols = [c for c in table_node.find_all("col") if c.name == "col"]
        if not cols:
            for cg in table_node.find_all("colgroup"):
                cols.extend([c for c in cg.find_all("col") if c.name == "col"])
        if len(cols) >= n_cols:
            raw = []
            for c in cols[:n_cols]:
                parsed = None
                for src in (c.get("style", ""), c.get("width", "")):
                    parsed = _parse_css_width(src)
                    if parsed is not None:
                        break
                raw.append(parsed if parsed is not None else 0.0)
            mm_vals = []
            for w in raw:
                if w is None or w <= 0:
                    mm_vals.append(0.0)
                elif w < 1.0:
                    mm_vals.append(w * total_mm)
                else:
                    mm_vals.append(w)
            fixed_sum = sum(mm_vals)
            if fixed_sum > 0:
                return [w * total_mm / fixed_sum for w in mm_vals]
        return [total_mm / n_cols] * n_cols

    # 优先从 <col> 读取列宽，否则按公式/文本权重均分
    col_widths_mm = _table_col_widths(table_node, page_width, n_cols)

    # 计算每列权重：若列内 cell 多数为纯公式，则给窄列；否则宽列
    col_weights = [3] * n_cols
    for ci in range(n_cols):
        pure_count = 0
        total = 0
        for row in logical:
            if ci < len(row):
                cell = row[ci]
                if cell is None or cell.get("_span"):
                    continue
                total += 1
                if _cell_pure_math(cell.get("node", None)):
                    pure_count += 1
        # 单列纯公式不再用 0.6 极窄权重，否则公式会被挤得很大；改用 1.2
        if total > 0 and pure_count == total:
            col_weights[ci] = 1.2
        elif total > 0 and pure_count / total >= 0.8:
            col_weights[ci] = 1.2

    # 未解析到 <col> 宽度时，使用权重分配
    if all(abs(w - col_widths_mm[0]) < 0.01 for w in col_widths_mm):
        total_weight = sum(col_weights)
        col_widths_mm = [page_width * w / total_weight for w in col_weights]

    with pdf.table(
        width=page_width,
        col_widths=tuple(col_widths_mm),
        text_align="LEFT",
        line_height=6,
        first_row_as_headings=False,
        borders_layout="ALL",
        cell_fill_mode="NONE",
    ) as _table:
        for ri, row in enumerate(logical):
            trow = _table.row()
            for ci in range(n_cols):
                cell = row[ci] if ci < len(row) else None
                if cell is None or cell.get("_span"):
                    continue
                cs = int(cell.get("cs", 1))
                rs = int(cell.get("rs", 1))
                node = cell.get("node")
                if not node:
                    trow.cell("", colspan=cs, rowspan=rs)
                    continue
                pure = _cell_pure_math(node)
                if pure:
                    tex, disp, png = pure
                    # 保持公式原始宽高比，按自然尺寸居中显示
                    trow.cell(img=str(png), colspan=cs, rowspan=rs,
                              img_fill_width=False,
                              align="CENTER",
                              v_align="MIDDLE")
                else:
                    # 无 math 的纯文本单元格直接写入文字，避免栅格化拉伸
                    clone = _copy.deepcopy(node)
                    if not clone.find("math"):
                        for br in list(clone.find_all("br")):
                            br.replace_with("\n")
                        for img in list(clone.find_all("img")):
                            alt = img.get("alt", "")
                            img.replace_with(f"[{alt or '图'}]")
                        txt = clone.get_text(" ", strip=True)
                        trow.cell(txt or "", colspan=cs, rowspan=rs)
                        continue
                    # 混合内容：按该列实际宽度渲染 PNG，不强制拉伸
                    col_w_mm = sum(col_widths_mm[ci:ci + cs]) if cs > 1 else col_widths_mm[ci]
                    col_w_px = int(col_w_mm / 25.4 * 200 * 0.92)
                    from paper_tools.core.exporter import _find_font, _CN_FONT_CANDIDATES
                    font_path = _find_font(_CN_FONT_CANDIDATES)
                    rendered_png = None
                    if font_path:
                        rendered_png = _render_mixed_cell_png(
                            node, font_path=font_path, target_w_px=col_w_px,
                            cache_dir=cache_dir,
                        )
                    if rendered_png:
                        trow.cell(img=str(rendered_png), colspan=cs,
                                  rowspan=rs, img_fill_width=False,
                                  align="CENTER", v_align="MIDDLE")
                    else:
                        # fallback：把公式退化为 unicode/简化文本
                        n2 = _copy.deepcopy(node)
                        for mn in list(n2.find_all("math")):
                            tex, disp = _extract_math_tex(mn)
                            if tex:
                                short = _math_to_unicode(tex)
                                mn.replace_with(f" {short} ")
                            else:
                                mn.decompose()
                        for img in list(n2.find_all("img")):
                            src = img.get("src", "")
                            alt = img.get("alt", "")
                            if src and not src.startswith("http"):
                                img.replace_with(f" [图] ")
                            elif alt:
                                img.replace_with(f" [{alt}] ")
                            else:
                                img.decompose()
                        txt = n2.get_text(" ", strip=True)
                        trow.cell(txt, colspan=cs, rowspan=rs)
    pdf.ln(3)


def _md_to_pdf_fallback(md_path: Path, pdf_path: Path) -> Path:
    """基于 Markdown 的降级 PDF 导出。"""
    from fpdf import FPDF

    cn_font = _find_font(_CN_FONT_CANDIDATES)
    if cn_font is None:
        raise FileNotFoundError("无法找到中文字体文件")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.add_font("CNBody", "", cn_font, uni=True)
    mono_font = _find_font(_MONO_FONT_CANDIDATES)
    if mono_font:
        pdf.add_font("CNMono", "", mono_font, uni=True)

    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.split("\n")
    n = len(lines)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin

    def _collect_table(start):
        buf = []
        j = start
        while j < n and _BLOCK_TABLE_ROW_RE.match(lines[j]):
            buf.append(lines[j]); j += 1
        return _parse_table_rows(buf), j

    h1_done = False
    i = 0
    while i < n:
        line = lines[i].rstrip(); i += 1
        if not line.strip():
            pdf.ln(3); continue
        if _BLOCK_HR_RE.match(line.strip()):
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y()); pdf.ln(5); continue
        hm = _BLOCK_HEADING_RE.match(line)
        if hm:
            level = len(hm.group(1)); sz = {1:18,2:15,3:13,4:12}.get(level,12)
            if level == 1 and not h1_done:
                h1_done = True
                pdf.ln(3); pdf.set_font("CNBody","B",sz); pdf.multi_cell(page_width, sz*0.7, _strip_inline(hm.group(2)), align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT); pdf.ln(6)
            elif level != 1:
                pdf.ln(3); pdf.set_font("CNBody","B",sz); pdf.multi_cell(page_width, sz*0.7, _strip_inline(hm.group(2)), new_x=XPos.LMARGIN, new_y=YPos.NEXT); pdf.ln(2)
            continue
        if _BLOCK_BLOCKQUOTE_RE.match(line):
            i -= 1
            while i < n and _BLOCK_BLOCKQUOTE_RE.match(lines[i]):
                pdf.set_font("CNBody","I",9); pdf.set_text_color(120,120,120)
                pdf.set_x(pdf.l_margin+8); pdf.multi_cell(page_width-8,6,_BLOCK_BLOCKQUOTE_RE.match(lines[i]).group(1), new_x=XPos.LMARGIN, new_y=YPos.NEXT); i+=1
            pdf.set_text_color(0,0,0); continue
        if _BLOCK_TABLE_ROW_RE.match(line):
            i -= 1; rows, ni = _collect_table(i); i = ni
            if rows:
                ncols = max(len(r) for r in rows)
                pdf.ln(2)
                for row_data in rows:
                    while len(row_data) < ncols: row_data.append("")
                    yb = pdf.get_y(); maxh = 0; xs = pdf.l_margin
                    for ci, ct in enumerate(row_data):
                        x = xs + ci*(page_width/ncols)
                        pdf.set_xy(x, yb); pdf.set_font("CNBody","B" if ci==0 else "",9)
                        t = _strip_inline(ct)
                        if len(t)>60: t = t[:57]+"..."
                        pdf.multi_cell(page_width/ncols,6,t,border=1,align="L")
                        maxh = max(maxh, pdf.get_y()-yb)
                    pdf.set_y(yb+maxh)
                pdf.ln(3)
            continue
        if line.startswith("$$") and line.endswith("$$"):
            pdf.ln(2); pdf.set_font("CNBody","I",10); pdf.set_text_color(90,90,90)
            pdf.multi_cell(page_width,7,line,align="C"); pdf.set_text_color(0,0,0); pdf.ln(2); continue
        pdf.set_font("CNBody","",11); pdf.multi_cell(page_width,7,_strip_inline(line)); pdf.ln(1)

    pdf.output(str(pdf_path))
    logger.info(f"PDF 已导出（基于 markdown 降级）: {pdf_path}")
    return pdf_path
